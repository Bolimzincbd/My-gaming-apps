from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "PROJECT_REPORT.md"
OUTPUT = ROOT / "docs" / "MLBB_Nexus_Project_Report.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_run_font(run, name="Aptos", size=10.5, bold=False, color="F4F7FB"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_code_block(doc, lines):
    if not lines:
        return
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, "111827")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run("\n")
        run = paragraph.add_run(line)
        set_run_font(run, name="Consolas", size=8.5, color="D7E3F0")
    doc.add_paragraph()


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    set_run_font(run, color="1F2937")


def add_paragraph(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(text)
    set_run_font(run, color="1F2937")


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(25)
    styles["Heading 1"].font.name = "Aptos Display"
    styles["Heading 1"].font.size = Pt(17)
    styles["Heading 2"].font.name = "Aptos Display"
    styles["Heading 2"].font.size = Pt(13)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("MLBB Nexus Project Report")
    set_run_font(title_run, name="Aptos Display", size=26, bold=True, color="111827")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Mobile Legends: Bang Bang Game Matcher and Marketplace Platform with Mock Escrow Workflow")
    set_run_font(subtitle_run, size=11, color="4B5563")

    note_table = doc.add_table(rows=1, cols=1)
    note_cell = note_table.cell(0, 0)
    set_cell_shading(note_cell, "FFF3CD")
    note_paragraph = note_cell.paragraphs[0]
    note_run = note_paragraph.add_run("Important: mock escrow is a university demo workflow only. It is not real payment escrow.")
    set_run_font(note_run, size=10.5, bold=True, color="5A4100")
    doc.add_paragraph()

    in_code = False
    code_lines = []
    for raw_line in SOURCE.read_text(encoding="utf-8").splitlines():
        line = raw_line.rstrip()
        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if not line:
            continue
        if line.startswith("# "):
            continue
        if line.startswith("## "):
            paragraph = doc.add_heading(line[3:], level=1)
            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.paragraph_format.space_after = Pt(4)
            continue
        if line.startswith("### "):
            paragraph = doc.add_heading(line[4:], level=2)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(2)
            continue
        if line.startswith("- "):
            add_bullet(doc, line[2:])
            continue
        add_paragraph(doc, line)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("MLBB Nexus | University Practicum Demo | Mock escrow only")
    set_run_font(footer_run, size=8.5, color="6B7280")

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
